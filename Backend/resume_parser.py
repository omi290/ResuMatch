# resume_parser.py - Enhanced version with RapidAPI integration and fallback parser
import io
import PyPDF2
import docx
import re
import logging
import json
from datetime import datetime
import os
import requests
from dotenv import load_dotenv
from thefuzz import fuzz

# Try to import config from different possible locations
try:
    from Backend.config import RAPIDAPI_KEY, RAPIDAPI_HOST, RAPIDAPI_URL, RAPIDAPI_HEADERS
except ImportError:
    try:
        from config import RAPIDAPI_KEY, RAPIDAPI_HOST, RAPIDAPI_URL, RAPIDAPI_HEADERS
    except ImportError:
        # Set default values if config is not available
        RAPIDAPI_KEY = os.getenv("RAPIDAPI_KEY")
        RAPIDAPI_HOST = "extracta.p.rapidapi.com"
        RAPIDAPI_URL = "https://extracta.p.rapidapi.com/extract"
        RAPIDAPI_HEADERS = {
            "X-RapidAPI-Key": RAPIDAPI_KEY,
            "X-RapidAPI-Host": RAPIDAPI_HOST
        } if RAPIDAPI_KEY else {}

# PyPDF2, docx, re, and other local parsing modules are no longer directly used for parsing
# as this functionality is replaced by RapidAPI.
# import PyPDF2
# import docx
# import re
# import random # Retained if used for other purposes, but not parsing logic

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

def extract_text_from_pdf(file_path):
    """Extracts text from a PDF file."""
    text = ""
    try:
        logger.info(f"Attempting to extract text from PDF: {file_path}")
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            num_pages = len(reader.pages)
            logger.info(f"PDF has {num_pages} pages")
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                page_text = page.extract_text() or ""
                text += page_text
                logger.debug(f"Extracted {len(page_text)} characters from page {page_num + 1}")
        logger.info(f"Successfully extracted {len(text)} total characters from PDF")
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}", exc_info=True)
    return text

def extract_text_from_docx(file_path):
    """Extracts text from a DOCX file."""
    text = ""
    try:
        logger.info(f"Attempting to extract text from DOCX: {file_path}")
        doc = docx.Document(file_path)
        num_paragraphs = len(doc.paragraphs)
        logger.info(f"DOCX has {num_paragraphs} paragraphs")
        for para in doc.paragraphs:
            para_text = para.text + "\n"
            text += para_text
            logger.debug(f"Extracted paragraph: {para_text[:100]}...")
        logger.info(f"Successfully extracted {len(text)} total characters from DOCX")
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}", exc_info=True)
    return text

def extract_text_from_file(file_path):
    """Determines file type and extracts text accordingly."""
    file_path_lower = file_path.lower()
    logger.info(f"Determining file type for: {file_path}")
    
    if file_path_lower.endswith('.pdf'):
        logger.info("File identified as PDF")
        return extract_text_from_pdf(file_path)
    elif file_path_lower.endswith(('.docx', '.doc')):
        logger.info("File identified as DOCX/DOC")
        return extract_text_from_docx(file_path)
    else:
        logger.warning(f"Unsupported file type for local parsing: {file_path}")
        return ""

def identify_sections(text):
    """Identifies common resume sections based on keywords."""
    sections = {}
    # Common section titles (can be expanded)
    section_titles = [
        "summary", "objective", "experience", "education", "skills",
        "awards", "certifications", "projects", "publications", "languages",
        "interests", "contact", "about", "work experience", "employment history",
        "professional experience", "technical skills", "core competencies",
        "key skills", "expertise", "qualifications", "achievements", "accomplishments",
        "volunteer", "community service", "extracurricular", "activities"
    ]
    
    # Create a regex pattern to find section titles
    # Case-insensitive, starts at beginning of line, optional leading whitespace
    pattern = re.compile(r"^(?:\s*|)([A-Z][a-zA-Z\s]+):?$", re.MULTILINE)
    
    # Additional patterns for skills sections
    skills_patterns = [
        r"^(?:\s*|)(technical\s+skills?|core\s+competencies?|key\s+skills?|expertise|qualifications?):?$",
        r"^(?:\s*|)(programming\s+languages?|languages?|technologies?|tools?|frameworks?):?$",
        r"^(?:\s*|)(software\s+skills?|computer\s+skills?|it\s+skills?):?$"
    ]

    lines = text.split('\n')
    current_section = "header" # Default section before first recognized title
    section_content = []

    for line in lines:
        line_stripped = line.strip()
        
        # Check for standard section titles
        match = pattern.match(line_stripped)
        if match:
            title = match.group(1).strip().lower()
            if title in section_titles:
                if current_section:
                    sections[current_section] = section_content
                current_section = title
                section_content = [line_stripped] # Include the section title line
                continue
        
        # Check for skills-related patterns
        for skills_pattern in skills_patterns:
            if re.match(skills_pattern, line_stripped, re.IGNORECASE):
                if current_section:
                    sections[current_section] = section_content
                current_section = "skills"
                section_content = [line_stripped]
                break
        else:
            # If no pattern matched, add to current section
            section_content.append(line_stripped)
    
    if current_section:
        sections[current_section] = section_content
    
    # Post-process to merge similar sections
    if "work experience" in sections and "experience" not in sections:
        sections["experience"] = sections["work experience"]
        del sections["work experience"]
    
    if "employment history" in sections and "experience" not in sections:
        sections["experience"] = sections["employment history"]
        del sections["employment history"]
    
    if "professional experience" in sections and "experience" not in sections:
        sections["experience"] = sections["professional experience"]
        del sections["professional experience"]
    
    # Merge technical skills, core competencies, etc. into skills section
    skills_related_sections = ["technical skills", "core competencies", "key skills", "expertise", "qualifications"]
    for section_name in skills_related_sections:
        if section_name in sections and "skills" not in sections:
            sections["skills"] = sections[section_name]
            del sections[section_name]
        elif section_name in sections and "skills" in sections:
            # Merge content
            sections["skills"].extend(sections[section_name])
            del sections[section_name]
    
    return sections

def extract_name_email(text_lines):
    """Extracts name and email from a list of text lines."""
    name = ""
    email = ""
    email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')

    for line in text_lines:
        # Extract email
        email_match = email_pattern.search(line)
        if email_match:
            email = email_match.group(0)
            
        # Basic name extraction (usually the first prominent line not containing email)
        if not name and not email_match and len(line.strip().split()) <= 4 and line.strip() and '@' not in line:
            name = line.strip()
            
    # Try to extract name from email if not found
    if not name and email:
        try:
            name = email.split('@')[0].replace('.', ' ').replace('_', ' ').title()
        except:
            pass

    return name, email

def extract_skills(text, skills_section_lines):
    """Extracts skills strictly from the skills section lines, using normalization and deduplication. Special handling for C++."""
    skills = set()

    def create_skill_variations(skill_name):
        """Create common variations of a skill name for better matching."""
        variations = [skill_name.lower()]
        
        # Handle common programming language variations
        if skill_name.lower() == "javascript":
            variations.extend(["js", "javascript", "ecmascript", "es6", "es2015"])
        elif skill_name.lower() == "typescript":
            variations.extend(["ts", "typescript"])
        elif skill_name.lower() == "c++":
            variations.extend(["cpp", "c plus plus", "c++", "c++ programming", "c plus plus programming"])
        elif skill_name.lower() == "c#":
            variations.extend(["csharp", "c sharp", "c# programming", "dotnet", ".net"])
        elif skill_name.lower() == "c":
            variations.extend(["c programming", "c language", "ansi c"])
        elif skill_name.lower() == "java":
            variations.extend(["java programming", "java language", "j2ee", "j2se", "jdk"])
        elif skill_name.lower() == "python":
            variations.extend(["python programming", "python language", "py"])
        elif skill_name.lower() == "php":
            variations.extend(["php programming", "php language"])
        elif skill_name.lower() == "ruby":
            variations.extend(["ruby programming", "ruby language", "rails", "ruby on rails"])
        elif skill_name.lower() == "go":
            variations.extend(["golang", "go programming", "go language"])
        elif skill_name.lower() == "rust":
            variations.extend(["rust programming", "rust language"])
        elif skill_name.lower() == "swift":
            variations.extend(["swift programming", "swift language", "ios development"])
        elif skill_name.lower() == "kotlin":
            variations.extend(["kotlin programming", "kotlin language", "android development"])
        elif skill_name.lower() == "scala":
            variations.extend(["scala programming", "scala language"])
        elif skill_name.lower() == "perl":
            variations.extend(["perl programming", "perl language"])
        elif skill_name.lower() == "r":
            variations.extend(["r programming", "r language", "r studio", "rstudio"])
        elif skill_name.lower() == "matlab":
            variations.extend(["matlab programming", "matlab language"])
        elif skill_name.lower() == "assembly":
            variations.extend(["assembly language", "asm", "x86", "x64", "arm assembly"])
        elif skill_name.lower() == "shell scripting":
            variations.extend(["bash", "shell", "bash scripting", "shell script", "linux shell"])
        elif skill_name.lower() == "bash":
            variations.extend(["bash scripting", "bash shell", "linux bash"])
        elif skill_name.lower() == "powershell":
            variations.extend(["powershell scripting", "ps", "windows powershell"])
            
        # Handle database variations
        elif skill_name.lower() == "mysql":
            variations.extend(["mysql", "my sql", "mysql database", "mysql server"])
        elif skill_name.lower() == "postgresql":
            variations.extend(["postgres", "postgresql", "postgres database", "postgresql database"])
        elif skill_name.lower() == "mongodb":
            variations.extend(["mongo", "mongodb", "mongo db", "mongo database"])
        elif skill_name.lower() == "sql":
            variations.extend(["sql", "structured query language", "sql database", "sql server"])
        elif skill_name.lower() == "nosql":
            variations.extend(["nosql", "no sql", "non-relational database"])
        elif skill_name.lower() == "oracle":
            variations.extend(["oracle database", "oracle db", "oracle sql"])
        elif skill_name.lower() == "sqlite":
            variations.extend(["sqlite", "sqlite database", "sqlite db"])
        elif skill_name.lower() == "sql server":
            variations.extend(["mssql", "microsoft sql server", "sql server database"])
        elif skill_name.lower() == "redis":
            variations.extend(["redis", "redis database", "redis cache"])
        elif skill_name.lower() == "cassandra":
            variations.extend(["cassandra", "apache cassandra", "cassandra database"])
        elif skill_name.lower() == "dynamodb":
            variations.extend(["dynamodb", "dynamo db", "aws dynamodb"])
        elif skill_name.lower() == "firebase":
            variations.extend(["firebase", "firebase database", "google firebase"])
        elif skill_name.lower() == "firestore":
            variations.extend(["firestore", "firestore database", "google firestore"])
        elif skill_name.lower() == "dbms":
            variations.extend(["dbms", "database management system", "database management"])
            
        # Handle web technology variations
        elif skill_name.lower() == "react":
            variations.extend(["react", "reactjs", "react.js", "react js", "facebook react"])
        elif skill_name.lower() == "angular":
            variations.extend(["angular", "angularjs", "angular.js", "angular js"])
        elif skill_name.lower() == "vue.js":
            variations.extend(["vue", "vue.js", "vue js", "vuejs"])
        elif skill_name.lower() == "vue":
            variations.extend(["vue", "vue.js", "vue js", "vuejs"])
        elif skill_name.lower() == "jquery":
            variations.extend(["jquery", "jquery.js", "jquery js"])
        elif skill_name.lower() == "bootstrap":
            variations.extend(["bootstrap", "twitter bootstrap", "bootstrap css"])
        elif skill_name.lower() == "tailwind css":
            variations.extend(["tailwind", "tailwind css", "tailwindcss"])
        elif skill_name.lower() == "node.js":
            variations.extend(["node", "node.js", "node js", "nodejs"])
        elif skill_name.lower() == "node":
            variations.extend(["node", "node.js", "node js", "nodejs"])
        elif skill_name.lower() == "express.js":
            variations.extend(["express", "express.js", "express js", "expressjs"])
        elif skill_name.lower() == "express":
            variations.extend(["express", "express.js", "express js", "expressjs"])
        elif skill_name.lower() == "django":
            variations.extend(["django", "django framework", "python django"])
        elif skill_name.lower() == "flask":
            variations.extend(["flask", "flask framework", "python flask"])
        elif skill_name.lower() == "laravel":
            variations.extend(["laravel", "laravel framework", "php laravel"])
        elif skill_name.lower() == "asp.net":
            variations.extend(["asp.net", "aspnet", "asp .net", "dotnet", ".net"])
        elif skill_name.lower() == "spring":
            variations.extend(["spring", "spring framework", "java spring"])
        elif skill_name.lower() == "spring boot":
            variations.extend(["spring boot", "springboot", "spring-boot"])
        elif skill_name.lower() == "hibernate":
            variations.extend(["hibernate", "hibernate orm", "java hibernate"])
        elif skill_name.lower() == "jpa":
            variations.extend(["jpa", "java persistence api", "jpa hibernate"])
            
        # Handle version control variations
        elif skill_name.lower() == "git":
            variations.extend(["git", "git version control", "git vcs"])
        elif skill_name.lower() == "github":
            variations.extend(["github", "git hub", "github.com"])
        elif skill_name.lower() == "gitlab":
            variations.extend(["gitlab", "git lab", "gitlab.com"])
        elif skill_name.lower() == "bitbucket":
            variations.extend(["bitbucket", "bit bucket", "bitbucket.org"])
        elif skill_name.lower() == "svn":
            variations.extend(["svn", "subversion", "apache svn"])
        elif skill_name.lower() == "subversion":
            variations.extend(["svn", "subversion", "apache svn"])
            
        # Handle DevOps variations
        elif skill_name.lower() == "docker":
            variations.extend(["docker", "docker container", "dockerization"])
        elif skill_name.lower() == "kubernetes":
            variations.extend(["kubernetes", "k8s", "kube", "container orchestration"])
        elif skill_name.lower() == "jenkins":
            variations.extend(["jenkins", "jenkins ci", "jenkins pipeline"])
        elif skill_name.lower() == "ci/cd":
            variations.extend(["ci/cd", "cicd", "continuous integration", "continuous deployment"])
        elif skill_name.lower() == "continuous integration":
            variations.extend(["ci/cd", "cicd", "continuous integration", "continuous deployment"])
        elif skill_name.lower() == "aws":
            variations.extend(["aws", "amazon web services", "amazon aws"])
        elif skill_name.lower() == "azure":
            variations.extend(["azure", "microsoft azure", "azure cloud"])
        elif skill_name.lower() == "google cloud platform":
            variations.extend(["gcp", "google cloud", "google cloud platform", "google cloud platform"])
        elif skill_name.lower() == "gcp":
            variations.extend(["gcp", "google cloud", "google cloud platform", "google cloud platform"])
        elif skill_name.lower() == "cloud computing":
            variations.extend(["cloud", "cloud computing", "cloud platform"])
            
        # Handle ML/AI variations
        elif skill_name.lower() == "machine learning":
            variations.extend(["ml", "machine learning", "machinelearning", "ml algorithms"])
        elif skill_name.lower() == "ml":
            variations.extend(["ml", "machine learning", "machinelearning", "ml algorithms"])
        elif skill_name.lower() == "deep learning":
            variations.extend(["deep learning", "deeplearning", "neural networks", "neural network"])
        elif skill_name.lower() == "ai":
            variations.extend(["ai", "artificial intelligence", "artificialintelligence"])
        elif skill_name.lower() == "artificial intelligence":
            variations.extend(["ai", "artificial intelligence", "artificialintelligence"])
        elif skill_name.lower() == "tensorflow":
            variations.extend(["tensorflow", "tf", "google tensorflow"])
        elif skill_name.lower() == "pytorch":
            variations.extend(["pytorch", "torch", "facebook pytorch"])
        elif skill_name.lower() == "scikit-learn":
            variations.extend(["scikit-learn", "sklearn", "scikit learn"])
        elif skill_name.lower() == "keras":
            variations.extend(["keras", "keras framework"])
        elif skill_name.lower() == "opencv":
            variations.extend(["opencv", "open cv", "computer vision library"])
        elif skill_name.lower() == "nltk":
            variations.extend(["nltk", "natural language toolkit"])
        elif skill_name.lower() == "computer vision":
            variations.extend(["computer vision", "cv", "image processing"])
        elif skill_name.lower() == "natural language processing":
            variations.extend(["nlp", "natural language processing", "text processing"])
        elif skill_name.lower() == "nlp":
            variations.extend(["nlp", "natural language processing", "text processing"])
        elif skill_name.lower() == "data science":
            variations.extend(["data science", "datascience", "data scientist"])
            
        # Handle software development variations
        elif skill_name.lower() == "data structures":
            variations.extend(["data structures", "data structure", "ds"])
        elif skill_name.lower() == "algorithms":
            variations.extend(["algorithms", "algorithm", "algo"])
        elif skill_name.lower() == "dsa":
            variations.extend(["dsa", "data structures and algorithms", "data structures algorithms"])
        elif skill_name.lower() == "object-oriented programming":
            variations.extend(["oop", "object oriented programming", "object-oriented programming"])
        elif skill_name.lower() == "oop":
            variations.extend(["oop", "object oriented programming", "object-oriented programming"])
        elif skill_name.lower() == "design patterns":
            variations.extend(["design patterns", "design pattern", "software patterns"])
        elif skill_name.lower() == "software architecture":
            variations.extend(["software architecture", "system architecture", "architecture"])
        elif skill_name.lower() == "microservices":
            variations.extend(["microservices", "microservice", "micro service", "micro services"])
        elif skill_name.lower() == "restful apis":
            variations.extend(["rest", "restful", "rest api", "restful api", "restful apis"])
        elif skill_name.lower() == "rest api":
            variations.extend(["rest", "restful", "rest api", "restful api", "restful apis"])
        elif skill_name.lower() == "graphql":
            variations.extend(["graphql", "graph ql", "graph-ql"])
        elif skill_name.lower() == "soap":
            variations.extend(["soap", "soap api", "soap web service"])
        elif skill_name.lower() == "web services":
            variations.extend(["web services", "web service", "api services"])
        elif skill_name.lower() == "api development":
            variations.extend(["api development", "api dev", "api programming"])
            
        # Handle development methodologies
        elif skill_name.lower() == "agile":
            variations.extend(["agile", "agile methodology", "agile development"])
        elif skill_name.lower() == "scrum":
            variations.extend(["scrum", "scrum methodology", "scrum master"])
        elif skill_name.lower() == "kanban":
            variations.extend(["kanban", "kanban board", "kanban methodology"])
        elif skill_name.lower() == "devops":
            variations.extend(["devops", "dev ops", "development operations"])
        elif skill_name.lower() == "tdd":
            variations.extend(["tdd", "test driven development", "test-driven development"])
        elif skill_name.lower() == "test-driven development":
            variations.extend(["tdd", "test driven development", "test-driven development"])
        elif skill_name.lower() == "bdd":
            variations.extend(["bdd", "behavior driven development", "behavior-driven development"])
        elif skill_name.lower() == "behavior-driven development":
            variations.extend(["bdd", "behavior driven development", "behavior-driven development"])
        elif skill_name.lower() == "waterfall":
            variations.extend(["waterfall", "waterfall methodology", "waterfall model"])
        elif skill_name.lower() == "lean development":
            variations.extend(["lean", "lean development", "lean methodology"])
            
        return variations
    
    # Expanded skills database from user, organized by categories
    skills_database = {
        # Programming and Development
        "programming_and_development": [
            # Core Programming Languages
            "Python", "Java", "C", "C++", "C#", "JavaScript", "JS", "HTML", "CSS", "TypeScript", "TS",
            "PHP", "Ruby", "Go", "Rust", "Swift", "Kotlin", "Scala", "Perl", "R", "MATLAB",
            "Assembly", "Shell Scripting", "Bash", "PowerShell",
            
            # Databases and Data Management
            "SQL", "MySQL", "PostgreSQL", "MongoDB", "NoSQL", "Oracle", "SQLite", "SQL Server",
            "DBMS", "Database Management", "Data Modeling", "ERD", "Entity Relationship",
            "Redis", "Cassandra", "DynamoDB", "Firebase", "Firestore",
            
            # Web Technologies
            "React", "Angular", "Vue.js", "Vue", "jQuery", "Bootstrap", "Tailwind CSS",
            "Node.js", "Node", "Express.js", "Express", "Django", "Flask", "Laravel",
            "ASP.NET", "Spring", "Spring Boot", "Hibernate", "JPA",
            
            # Version Control and DevOps
            "Git", "GitHub", "GitLab", "Bitbucket", "SVN", "Subversion",
            "Docker", "Kubernetes", "Jenkins", "CI/CD", "Continuous Integration",
            "AWS", "Azure", "Google Cloud Platform", "GCP", "Cloud Computing",
            
            # Machine Learning and AI
            "Machine Learning", "ML", "Deep Learning", "AI", "Artificial Intelligence",
            "TensorFlow", "PyTorch", "Scikit-learn", "Keras", "OpenCV", "NLTK",
            "Computer Vision", "Natural Language Processing", "NLP", "Data Science",
            
            # Software Development
            "Data Structures", "Algorithms", "DSA", "Object-Oriented Programming", "OOP",
            "Design Patterns", "Software Architecture", "Microservices", "RESTful APIs", "REST API",
            "GraphQL", "SOAP", "Web Services", "API Development",
            
            # Development Methodologies
            "Agile", "Scrum", "Kanban", "DevOps", "TDD", "Test-Driven Development",
            "BDD", "Behavior-Driven Development", "Waterfall", "Lean Development"
        ],
        # Data and Analytics
        "data_and_analytics": [
            "Microsoft Excel", "Excel", "Power BI", "Tableau", "QlikView", "Qlik Sense",
            "Pandas", "NumPy", "Matplotlib", "Seaborn", "Plotly", "Bokeh",
            "R", "SAS", "SPSS", "Stata", "SQL for Analysis", "Data Analysis",
            "Data Visualization", "Big Data", "Hadoop", "Spark", "Apache Spark",
            "ETL", "Data Pipeline", "Statistical Modeling", "Predictive Analytics",
            "A/B Testing", "Business Intelligence", "BI", "Data Mining", "Data Warehousing"
        ],
        # Design and Multimedia
        "design_and_multimedia": [
            "Adobe Photoshop", "Photoshop", "Adobe Illustrator", "Illustrator", "Adobe InDesign", "InDesign",
            "Figma", "Sketch", "Adobe XD", "XD", "InVision", "Zeplin",
            "UI/UX Design", "User Interface", "User Experience", "Wireframing", "Prototyping",
            "Video Editing", "Animation", "Blender", "After Effects", "Premiere Pro",
            "3D Modeling", "Graphic Design", "Logo Design", "Branding"
        ],
        # Digital Marketing
        "digital_marketing": [
            "SEO", "Search Engine Optimization", "SEM", "Search Engine Marketing",
            "Google Ads", "Facebook Ads", "Social Media Advertising", "PPC", "Pay-Per-Click",
            "Google Analytics", "Google Tag Manager", "Email Marketing", "Mailchimp", "Sendinblue",
            "Content Marketing", "Social Media Management", "Copywriting", "Digital Marketing",
            "Marketing Automation", "Lead Generation", "Conversion Optimization"
        ],
        # Engineering and Technical
        "engineering_and_technical": [
            "AutoCAD", "SolidWorks", "MATLAB", "Simulink", "ANSYS", "CATIA",
            "PLC Programming", "Embedded Systems", "Microcontrollers", "Arduino", "Raspberry Pi",
            "Circuit Design", "PCB Design", "Robotics", "IoT", "Internet of Things",
            "Electrical Engineering", "Mechanical Engineering", "Civil Engineering",
            "Chemical Engineering", "Industrial Engineering"
        ],
        # Cognitive Skills
        "cognitive_skills": [
            "Critical Thinking", "Problem Solving", "Analytical Thinking", "Logical Thinking",
            "Decision Making", "Creativity", "Attention to Detail", "Strategic Thinking",
            "Innovation", "Research Skills", "Data Analysis", "Quantitative Analysis"
        ],
        # Communication Skills
        "communication_skills": [
            "Verbal Communication", "Written Communication", "Public Speaking", "Presentation Skills",
            "Active Listening", "Conflict Resolution", "Empathy", "Interpersonal Skills",
            "Negotiation", "Persuasion", "Technical Writing", "Documentation"
        ],
        # Team and Leadership
        "team_and_leadership": [
            "Teamwork", "Collaboration", "Leadership", "Team Management", "Project Management",
            "Delegation", "Mentoring", "Coaching", "Stakeholder Management", "Cross-functional Teams",
            "Agile Leadership", "Scrum Master", "Product Owner", "Program Management"
        ],
        # Organizational Skills
        "organizational_skills": [
            "Time Management", "Multitasking", "Goal Setting", "Prioritization", "Adaptability",
            "Stress Management", "Planning", "Organization", "Efficiency", "Productivity",
            "Resource Management", "Risk Management"
        ],
        # Language Skills
        "language_skills": [
            "English", "Hindi", "Spanish", "French", "German", "Chinese", "Japanese", "Korean",
            "Arabic", "Portuguese", "Italian", "Russian", "Translation", "Interpretation",
            "Bilingual", "Multilingual", "Language Proficiency"
        ],
        # Other Professional Skills
        "other_professional_skills": [
            "Resume Writing", "Business Intelligence", "CRM", "Salesforce", "HubSpot", "Zoho",
            "Customer Service", "Technical Writing", "Legal Research", "Financial Modeling",
            "Budgeting", "Forecasting", "Supply Chain Management", "Inventory Management",
            "Quality Assurance", "QA", "Testing", "Manual Testing", "Automated Testing",
            "Selenium", "JUnit", "TestNG", "Cucumber", "JIRA", "Confluence", "Slack", "Microsoft Teams"
        ]
    }
    
    # Flatten all skills into one list for easier searching
    all_skills = []
    for category, skill_list in skills_database.items():
        all_skills.extend(skill_list)

    # Only use the skills section lines for extraction
    skills_text = " ".join(skills_section_lines).lower()
    found_skills = set()

    # Special handling for C++ (and variations)
    cpp_patterns = [r'\bc\+\+\b', r'\bcpp\b', r'\bc plus plus\b']
    for line in skills_section_lines:
        line_lower = line.lower()
        for pattern in cpp_patterns:
            if re.search(pattern, line_lower):
                found_skills.add("C++")
                break

    # Tokenize the skills section into words/phrases (including multi-word tokens)
    tokens = set(re.findall(r'\b\w[\w\s\-\.\/#\+]+\b', skills_text))
    individual_words = set(re.findall(r'\b\w+\b', skills_text))
    tokens.update(individual_words)

    for skill in all_skills:
        if skill == "C++":
            continue  # Already handled above
        skill_lower = skill.lower()
        skill_variations = create_skill_variations(skill)
        for variation in skill_variations:
            # Strict match only in the skills section
            if re.search(r'\b' + re.escape(variation) + r'\b', skills_text):
                found_skills.add(skill)
                break

    skills_list = sorted(list(found_skills), key=lambda x: x.lower())
    logger.info(f"Extracted {len(skills_list)} skills (strictly from skills section, with C++ fix): {skills_list}")
    return skills_list

def parse_experience_section(experience_lines):
    """Parses experience entries from a list of lines."""
    experiences = []
    current_exp = {}
    
    # Regex for date ranges (e.g., "Jan 2020 - Dec 2022", "2020 - Present")
    date_pattern = re.compile(r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}|\d{4})\s*[-–]\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}|\d{4}|Present)', re.IGNORECASE)
    
    # Keywords indicating a new experience entry or a detail line
    new_entry_keywords = ["company", "title", "role"]

    for line in experience_lines:
        line = line.strip()
        if not line:
            continue
        
        # Check for date range to identify new entries or duration
        date_match = date_pattern.search(line)

        # Check if line indicates a new experience entry (e.g., contains a common company/title pattern)
        is_new_entry = False
        if date_match:
            is_new_entry = True
        else:
            # Heuristic: If line looks like a title/company (e.g., "Software Engineer at Google")
            if re.search(r'\b(engineer|developer|manager|specialist|analyst|associate)\b', line, re.IGNORECASE) and len(line.split()) < 10:
                is_new_entry = True
            
        if is_new_entry:
            if current_exp:
                experiences.append(current_exp)
            current_exp = {"title": "", "company": "", "start_date": "", "end_date": "", "description": []}
            
            # Extract dates if found
            if date_match:
                current_exp["start_date"] = date_match.group(1).strip()
                current_exp["end_date"] = date_match.group(2).strip()
                line = date_pattern.sub("", line).strip() # Remove dates from the line
            
            # Try to extract title and company from the rest of the line
            if " at " in line:
                parts = line.split(" at ", 1)
                current_exp["title"] = parts[0].strip()
                current_exp["company"] = parts[1].strip()
            elif "," in line and len(line.split(',')) <= 3: # Simple comma separation
                parts = [p.strip() for p in line.split(',', 1)]
                current_exp["title"] = parts[0]
                if len(parts) > 1: current_exp["company"] = parts[1]
            else: # Fallback to assigning to title
                current_exp["title"] = line
        else:
            if current_exp and line:
                current_exp["description"].append(line)
    
    if current_exp:
        experiences.append(current_exp)
        
    # Join description lines into a single string
    for exp in experiences:
        exp["description"] = "\n".join(exp["description"]).strip()

    return experiences

def parse_education_section(education_lines):
    """Parses education entries from a list of lines, more robustly."""
    education_entries = []
    current_entry_lines = []

    # Regex for year ranges (e.g., "2019 - 2023", "2023", "Jan 2020 - Present")
    year_pattern = re.compile(
        r"(?:(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+)?(\d{4})"
        r"\s*[-–]?\s*"
        r"(?:(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+)?(\d{4}|Present|Current)?",
        re.IGNORECASE
    )

    # Keywords for degrees and institutions to help identify start of new entries
    degree_keywords = ["b.tech", "m.tech", "phd", "bachelor", "master", "degree", "diploma", "associate"]
    institution_keywords = ["university", "college", "institute", "school", "acad", "polytechnic"]

    def extract_edu_details(lines):
        degree = ""
        institution = ""
        year = ""
        description_lines = []
        text_block = " ".join(lines).strip()

        # Try to find year first
        year_match = year_pattern.search(text_block)
        if year_match:
            start_year = year_match.group(2) # The first 4-digit year
            end_year = year_match.group(4) # The second 4-digit year or 'Present'/'Current'
            
            if end_year:
                year = f"{start_year}-{end_year}"
            else:
                year = start_year
                
            # If 'Present' or 'Current' is in the original text, append it
            if re.search(r'Present|Current', text_block, re.IGNORECASE):
                if not end_year or end_year != 'Present': # Avoid double 'Present'
                    year = f"{start_year}-Present"

            # Remove year string from text_block for further parsing
            text_block = year_pattern.sub("", text_block, 1).strip() # Use count=1 to remove only the first match

        # Try to find degree and institution based on keywords/patterns
        for line in lines:
            lower_line = line.lower()
            
            # Degree extraction (more specific patterns)
            degree_match = re.search(r'\b(b\.tech|m\.tech|phd|bachelor.*?degree|master.*?degree|diploma|associate.*?degree)\b', lower_line)
            if degree_match and not degree:
                degree = re.sub(r'\b(in|of)\s+', '', degree_match.group(0), flags=re.IGNORECASE).strip()
                if not degree.endswith('degree') and not degree.endswith('tech') and not degree.endswith('ma') and not degree.endswith('phd'):
                    # Heuristic: if it's just a common word, make sure it's part of a degree phrase
                    if len(degree.split()) < 3 and not any(kw in degree for kw in ['engg', 'sci', 'arts']):
                        degree = ""

            # Institution extraction (look for common institution names or words)
            inst_match = re.search(r'\b([a-zA-Z\s&.-]+?(university|college|institute|school|acad|polytechnic|technical center))\b', lower_line)
            if inst_match and not institution:
                institution = inst_match.group(0).strip()
            
            # If no clear degree or institution yet, add to description
            if not (degree or institution) and line.strip() and not year_pattern.search(line):
                description_lines.append(line.strip())

        # Fallback if degree or institution still not found clearly
        if not degree and lines: # If first line might be the degree
            if len(lines[0].split()) < 7 and not any(kw in lines[0].lower() for kw in institution_keywords): # Short line, no institution kw
                degree = lines[0].strip()
                description_lines = lines[1:]
        
        if not institution and lines:
            # Try to get institution from second line if first was degree and second isn't too long
            if len(lines) > 1 and len(lines[1].split()) < 10 and not any(kw in lines[1].lower() for kw in degree_keywords):
                 institution = lines[1].strip()
                 if not degree: # If degree not set yet, assume first line was institution
                     degree = lines[0].strip()
                 description_lines = lines[2:]
            elif len(lines) == 1 and not degree: # If only one line and no degree, assume institution
                institution = lines[0].strip()
        
        # Ensure we don't have empty degree/institution if there's content
        if not degree and not institution and len(lines) > 0:
            if len(lines[0].split()) < 5: # Small chance it's a very short degree/institution
                if not any(k in lines[0].lower() for k in ['class', 'grade', 'marks']):
                    degree = lines[0].strip()
            elif len(lines) > 1 and len(lines[1].split()) < 5: # Or the second line is short
                 if not any(k in lines[1].lower() for k in ['class', 'grade', 'marks']):
                    institution = lines[1].strip()


        # Combine descriptions, filtering out common resume fluff (like percentages etc. already handled)
        description = " ".join(description_lines).strip()
        description = re.sub(r'\s{2,}', ' ', description).strip() # Reduce multiple spaces

        return {
            "degree": degree or "N/A",
            "institution": institution or "N/A",
            "year": year or "N/A",
            "description": description
        }


    for line in education_lines:
        line = line.strip()
        if not line:
            continue

        # Heuristic to detect a new education entry:
        # A line that looks like a degree, institution, or contains a year, and is relatively short.
        is_new_entry_candidate = False
        lower_line = line.lower()

        if year_pattern.search(line) and len(line.split()) < 15: # Year with short line
            is_new_entry_candidate = True
        elif any(kw in lower_line for kw in degree_keywords) and len(line.split()) < 15: # Degree keyword with short line
            is_new_entry_candidate = True
        elif any(kw in lower_line for kw in institution_keywords) and len(line.split()) < 15: # Institution keyword with short line
            is_new_entry_candidate = True

        # Also consider lines that are capitalized and look like titles/headings
        if line and line[0].isupper() and (not current_entry_lines or len(line.split()) < 5 or line.endswith(':')):
            is_new_entry_candidate = True


        if is_new_entry_candidate and current_entry_lines: # If it's a new entry AND we have collected previous lines
            education_entries.append(extract_edu_details(current_entry_lines))
            current_entry_lines = [line]
        elif not current_entry_lines and is_new_entry_candidate: # First entry is also a new entry candidate
             current_entry_lines = [line]
        else: # Otherwise, add to current entry
            current_entry_lines.append(line)

    if current_entry_lines:
        education_entries.append(extract_edu_details(current_entry_lines))
    
    return education_entries

def extract_extracurricular(extracurricular_lines):
    """Extracts extracurricular activities from a list of lines."""
    activities = []
    for line in extracurricular_lines:
        line = line.strip()
        # Filter out common section titles that might be accidentally included
        if line and not line.lower().startswith(("extracurricular activities", "hobbies", "interests", "awards", "certifications")):
            activities.append(line)
    return activities

def parse_resume_with_extracta(file_path):
    """
    Parse resume using the Extracta API from RapidAPI with fallback to improved basic parser
    
    Args:
        file_path (str): Path to the resume file
        
    Returns:
        dict: Parsed resume data
    """
    # Check if API key exists
    if not RAPIDAPI_KEY:
        logger.warning("No API key found. Using fallback parser.")
        return basic_resume_parser(file_path)
    
    try:
        with open(file_path, "rb") as file:
            files = {"file": file}
            logger.info(f"Sending request to parse resume: {os.path.basename(file_path)}")
            
            # Try with API first
            try:
                response = requests.post(RAPIDAPI_URL, files=files, headers=RAPIDAPI_HEADERS, timeout=15)
                
                # Log the response status and headers for debugging
                logger.info(f"Response status: {response.status_code}")
                
                if response.status_code == 200:
                    try:
                        data = response.json()
                        logger.info("Successfully parsed resume with API")
                        
                        # Save the raw response for debugging
                        debug_file = os.path.join(os.path.dirname(file_path), "parsed_debug.json")
                        with open(debug_file, 'w') as f:
                            json.dump(data, f, indent=2)
                        
                        # Process API response to ensure it matches our expected format
                        if "data" in data and "document" in data["data"]:
                            doc = data["data"]["document"]
                            
                            # Extract any sections that might be misplaced
                            extracted_text = extract_text_from_file(file_path)
                            sections = identify_sections(extracted_text)
                            
                            # Initialize additional_sections if not present
                            if "additional_sections" not in doc:
                                doc["additional_sections"] = {}
                            
                            # Handle certifications
                            if "certifications" in sections:
                                certifications = extract_certifications(sections["certifications"])
                                if certifications:
                                    doc["additional_sections"]["certifications"] = certifications
                            
                            # Handle extracurricular activities
                            if "extracurricular" in sections:
                                extracurricular = []
                                for line in sections["extracurricular"]:
                                    if line and len(line) < 100:
                                        extracurricular.append(line.strip())
                                if extracurricular:
                                    doc["additional_sections"]["extracurricular"] = extracurricular
                            
                            # Ensure skills is a list
                            if "skills" in doc and not isinstance(doc["skills"], list):
                                doc["skills"] = [doc["skills"]] if doc["skills"] else []
                            
                            # Ensure experience is a list with proper structure
                            if "experience" in doc:
                                if not isinstance(doc["experience"], list):
                                    doc["experience"] = [doc["experience"]] if doc["experience"] else []
                                
                                # Clean up experience entries
                                for i, exp in enumerate(doc["experience"]):
                                    if not isinstance(exp, dict):
                                        doc["experience"][i] = {
                                            "title": str(exp),
                                            "company": "",
                                            "duration": ""
                                        }
                            
                            # Ensure education is a list with proper structure
                            if "education" in doc:
                                if not isinstance(doc["education"], list):
                                    doc["education"] = [doc["education"]] if doc["education"] else []
                                
                                # Clean up education entries
                                for i, edu in enumerate(doc["education"]):
                                    if not isinstance(edu, dict):
                                        doc["education"][i] = {
                                            "degree": str(edu),
                                            "school": "",
                                            "duration": ""
                                        }
                        
                        return data
                    except json.JSONDecodeError:
                        logger.error("Failed to decode JSON response, using fallback parser")
                        logger.error(f"Response content: {response.text[:500]}...")
                else:
                    logger.error(f"API request failed with status code {response.status_code}")
                    logger.error(f"Response content: {response.text[:500]}...")
            except requests.exceptions.RequestException as e:
                logger.error(f"API request error: {str(e)}")
            
            # If we reach here, API parsing failed, use fallback
            logger.warning("Using fallback resume parser")
            return basic_resume_parser(file_path)
    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        return basic_resume_parser(file_path)

def basic_resume_parser(file_path):
    """
    Basic resume parser that uses local text extraction and parsing
    
    Args:
        file_path (str): Path to the resume file
        
    Returns:
        dict: Parsed resume data
    """
    try:
        # Extract text from file
        text = extract_text_from_file(file_path)
        if not text:
            logger.error("Failed to extract text from file")
            return None

        # Identify sections
        sections = identify_sections(text)
        
        # Extract basic information
        name, email = extract_name_email(sections.get("header", []))
        
        # Parse experience
        experience = parse_experience_section(sections.get("experience", []))
        
        # Parse education
        education = parse_education_section(sections.get("education", []))
        
        # Extract skills
        skills = extract_skills(text, sections.get("skills", []))
        
        # Extract certifications
        certifications = extract_certifications(sections.get("certifications", []))
        
        # Extract extracurricular activities separately
        extracurricular_activities = extract_extracurricular(sections.get("extracurricular", []))

        # Compile results
        result = {
            "data": {
                "document": {
                    "name": name,
                    "email": email,
                    "experience": experience,
                    "education": education,
                    "skills": skills,
                    "additional_sections": {
                        "certifications": certifications,
                        "extracurricular": extracurricular_activities
                    }
                }
            }
        }
        
        return result
    except Exception as e:
        logger.error(f"Error in basic resume parser: {str(e)}")
        return None

def extract_certifications(certification_lines):
    """Extracts certification entries from a list of lines."""
    certifications = []
    for line in certification_lines:
        line = line.strip()
        if line and len(line) < 100:  # Basic validation
            certifications.append(line)
    return certifications

def test_parser_configuration():
    """Test function to verify parser configuration"""
    try:
        logger.info("Testing resume parser configuration...")
        
        # Check RapidAPI configuration
        if not RAPIDAPI_KEY:
            logger.error("RAPIDAPI_KEY is not set")
            return False, "RAPIDAPI_KEY is not set"
            
        logger.info(f"RAPIDAPI configuration loaded successfully")
        logger.info(f"Host: {RAPIDAPI_HOST}")
        logger.info(f"URL: {RAPIDAPI_URL}")
        
        return True, "Configuration is valid"
    except Exception as e:
        logger.error(f"Error testing parser configuration: {str(e)}")
        return False, str(e)

# All previous local parsing functions are removed below this line.
# They are no longer needed as RapidAPI handles the parsing.
# For reference, removed functions were:
# - extract_text_from_pdf
# - extract_text_from_docx
# - extract_text_from_file
# - identify_sections
# - extract_name_email
# - parse_experience_section
# - parse_education_section
# - extract_skills
# - extract_certifications
# - basic_resume_parser
# (These comments are just for clarity and will not be in the final code)